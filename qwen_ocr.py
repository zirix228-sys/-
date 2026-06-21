import logging
import base64
import io
import time
import re
import os
import sys
import math
from datetime import datetime
from typing import Dict, Any, List, Optional
from ollama import Client
from PIL import Image
from tqdm import tqdm
from docx import Document

logger = logging.getLogger(__name__)

# Основной модуль OCR, перевода и RAG-поиска
class QwenOCR:
    def __init__(self, model: str = "qwen3-vl:8b-instruct", retriever=None):
        self.model = model
        self.host = "http://127.0.0.1:11434"
        self.client = Client(host=self.host)
        self.retriever = retriever
        self._check_ollama()

    def _check_ollama(self):
        try:
            models = self.client.list()
            model_names = [m.get("name", "") for m in models.get("models", [])]
            if self.model not in model_names:
                logger.warning(f"Модель '{self.model}' не найдена. Доступные: {model_names}")
            logger.info(f"Ollama подключена. Всего моделей: {len(model_names)}")
        except Exception as e:
            if sys.platform == 'win32':
                raise RuntimeError("Не удалось подключиться к Ollama. Убедись, что приложение Ollama запущено (иконка в трее).") from e
            else:
                raise RuntimeError(f"Не удалось подключиться к Ollama: {e}") from e

    def _pdf_to_images(self, pdf_bytes: bytes, dpi: int = 200) -> List[Image.Image]:
        try:
            from pdf2image import convert_from_bytes
            logger.info(f"Конвертация PDF в изображения (DPI={dpi})...")
            
            # Адаптация для Windows: нужен путь к Poppler
            kwargs = {"dpi": dpi, "thread_count": 2}
            if sys.platform == 'win32':
                poppler_path = os.environ.get('POPPLER_PATH', r'C:\poppler\Library\bin')
                if not os.path.exists(poppler_path):
                    logger.warning(f"Папка Poppler не найдена: {poppler_path}. Ошибка конвертации PDF почти гарантирована.")
                kwargs["poppler_path"] = poppler_path
            
            images = convert_from_bytes(pdf_bytes, **kwargs)
            
            if not images:
                raise ValueError("PDF пустой или не содержит страниц")
            processed = []
            for img in images:
                if img.mode in ("RGBA", "LA", "P"):
                    bg = Image.new("RGB", img.size, (255, 255, 255))
                    if img.mode == "P":
                        img = img.convert("RGBA")
                    bg.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
                    img = bg
                elif img.mode != "RGB":
                    img = img.convert("RGB")
                processed.append(img)
            logger.info(f"Конвертировано страниц: {len(processed)}")
            return processed
        except ImportError:
            if sys.platform == 'win32':
                raise ImportError(
                    "Для работы с PDF на Windows нужно:\n"
                    "1. pip install pdf2image\n"
                    "2. Скачать Poppler: https://github.com/oschwartz10612/poppler-windows/releases/\n"
                    "3. Распаковать его в C:\\poppler (или задать переменную среды POPPLER_PATH)."
                )
            else:
                raise ImportError("Установи: pip install pdf2image & sudo apt install poppler-utils")
        except Exception as e:
            logger.error(f"Ошибка конвертации PDF: {e}")
            raise

    def _docx_to_text(self, file_bytes: bytes) -> str:
        """Извлекает текст из DOCX-файла."""
        document = Document(io.BytesIO(file_bytes))
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n".join(parts)

    def _estimate_pages_by_chars(self, text: str, chars_per_page: int = 1800) -> int:
        return max(1, math.ceil(len(text) / chars_per_page))
    
    def _split_text(self, text: str, max_chars: int = 3500) -> List[str]:
        """Делит длинный текст на части, чтобы модель не теряла куски документа."""
        paragraphs = text.splitlines()
        chunks = []
        current = ""

        for paragraph in paragraphs:
            if len(current) + len(paragraph) + 1 > max_chars:
                if current.strip():
                    chunks.append(current.strip())
                current = paragraph
            else:
                current += "\n" + paragraph

        if current.strip():
            chunks.append(current.strip())

        return chunks

    def _prepare_image(self, image: Image.Image, max_size: tuple = (1536, 1536)) -> str:
        if image.mode != "RGB":
            image = image.convert("RGB")
        img_copy = image.copy()
        img_copy.thumbnail(max_size, Image.Resampling.LANCZOS)
        buffered = io.BytesIO()
        img_copy.save(buffered, format="JPEG", quality=85, optimize=True, progressive=True)
        return base64.b64encode(buffered.getvalue()).decode("utf-8")

    def _parse_response(self, response: Any) -> str:
        text = ""
        if isinstance(response, dict):
            text = response.get("response", "") or response.get("message", {}).get("content", "")
        elif hasattr(response, "response"):
            text = response.response
        elif hasattr(response, "message"):
            text = getattr(response.message, "get", lambda k, d: d)("content", "")
        return text.strip() if text else ""

    # Формирование контекста из параллельного корпуса
    def _build_rag_context(self, query_text: str, source_lang: str, target_lang: str, k: int = 5) -> str:
        """Универсальный RAG-контекст: работает для обоих направлений"""
        if not self.retriever:
            return ""
        try:
            query = query_text[:150].strip()
            if not query:
                return ""
            matches = self.retriever.search_parallel(query, k=k)
            if not matches:
                return ""
            
            if source_lang == "ru" and target_lang == "en":
                lines = [f"\nGLOSSARY (Russian → English):"]
                for m in matches:
                    lines.append(f"- RU: {m.get('ru_text', '')}\n  EN: {m.get('en_text', '')}")
            elif source_lang == "en" and target_lang == "ru":
                lines = [f"\nГЛОССАРИЙ (английский → русский):"]
                for m in matches:
                    lines.append(f"- EN: {m.get('en_text', '')}\n  RU: {m.get('ru_text', '')}")
            else:
                lines = [f"\nREFERENCE TERMS:"]
                for m in matches:
                    lines.append(f"- {m.get('ru_text', '')} / {m.get('en_text', '')}")
            
            return "\n".join(lines) + "\n"
        except Exception as e:
            logger.warning(f"Ошибка RAG-контекста: {e}")
            return ""

    def _fix_broken_tables(self, text: str, target_lang: str) -> str:
        """Нормализует Markdown-таблицы: выравнивает колонки, склеивает разрывы, убирает артефакты."""
        lines = text.splitlines()
        result = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('|') and '|' in line[1:]:
                table_block = [line]
                i += 1
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_block.append(lines[i].strip())
                    i += 1
                
                if len(table_block) >= 2:
                    try:
                        col_count = max(2, max(row.count('|') - 1 for row in table_block))
                        fixed_table = []
                        has_separator = False
                        
                        for idx, row in enumerate(table_block):
                            parts = [p.strip() for p in row.split('|')]
                            if parts and not parts[0]: parts.pop(0)
                            if parts and not parts[-1]: parts.pop()
                            
                            # Проверка на разделитель
                            if not has_separator and all(re.fullmatch(r'[\s\-\:\|]*', p) for p in parts):
                                sep_cells = []
                                for ci in range(col_count):
                                    if ci < len(parts) and ':' in parts[ci]:
                                        if parts[ci].startswith(':') and parts[ci].endswith(':'):
                                            sep_cells.append(':---:')
                                        elif parts[ci].endswith(':'):
                                            sep_cells.append('---:')
                                        elif parts[ci].startswith(':'):
                                            sep_cells.append(':---')
                                        else:
                                            sep_cells.append('---')
                                    else:
                                        sep_cells.append('---')
                                fixed_table.append('| ' + ' | '.join(sep_cells) + ' |')
                                has_separator = True
                            else:
                                while len(parts) < col_count:
                                    parts.append('')
                                parts = parts[:col_count]
                                if target_lang == 'ru':
                                    parts = [p.replace('|', '&#124;') for p in parts]
                                fixed_table.append('| ' + ' | '.join(parts) + ' |')
                        
                        if not has_separator and len(fixed_table) > 1:
                            fixed_table.insert(1, '| ' + ' | '.join(['---'] * col_count) + ' |')
                        
                        result.extend(fixed_table)
                    except Exception:
                        result.extend(table_block)
                else:
                    result.extend(table_block)
                continue
            result.append(line)
            i += 1
        return '\n'.join(result)

    def _detect_source_language(self, preview_text: str) -> str:
        """Простое определение языка по первым символам"""
        ru_chars = sum(1 for c in preview_text if '\u0400' <= c <= '\u04FF')
        en_chars = sum(1 for c in preview_text if c.isalpha() and ord(c) < 128)
        
        if ru_chars > en_chars and ru_chars > 3:
            return "ru"
        return "en"

    # Подготовка инструкций для выбранного направления перевода
    def _get_prompts(self, source_lang: str, target_lang: str) -> tuple:
        """Возвращает (system_prompt, translation_instruction, src_code, tgt_code)"""
        if target_lang == "ru":
            src_code, tgt_code = "EN", "RU"
            translation_instruction = "Переведи КАЖДОЕ слово на РУССКИЙ язык. Используй термины из глоссария. Транслитерируй имена и адреса. НИ СЛОВА на английском!"
        else:  # target_lang == "en"
            src_code, tgt_code = "RU", "EN"
            translation_instruction = "Translate EVERY word to ENGLISH. Use glossary terms. Transliterate names and addresses. NO Russian words!"
        
        system_prompt = (
            f"Ты — экспертная система OCR и профессиональный переводчик банковских/юридических документов.\n"
            f"ТВОЯ ЕДИНСТВЕННАЯ ЗАДАЧА: извлечь текст с изображения и ВЫДАТЬ ЕГО ПОЛНЫЙ ПЕРЕВОД.\n\n"
            f"ИСХОДНЫЙ ЯЗЫК: {source_lang.upper()}\n"
            f"ЦЕЛЕВОЙ ЯЗЫК: {tgt_code}\n\n"
            f"❗ АБСОЛЮТНЫЕ ПРАВИЛА:\n"
            f"1. ЯЗЫК: {translation_instruction}\n"
            f"2. СТРУКТУРА: Сохраняй оригинальную структуру. Заголовки, абзацы, списки, нумерация — всё как в оригинале.\n"
            f"3. ТАБЛИЦЫ (КРИТИЧНО):\n"
            f"   - Если в оригинале таблица, выводи ТОЛЬКО строгий Markdown.\n"
            f"   - ЗАГОЛОВОК: | Кол1 | Кол2 | ... | КолN |\n"
            f"   - РАЗДЕЛИТЕЛЬ: |---|---|...|---| (строго один раз после заголовка)\n"
            f"   - КАЖДАЯ строка данных должна содержать РОВНО N колонок. Никогда не добавляй и не удаляй колонки.\n"
            f"   - Если текст в ячейке многострочный, объединяй его в ОДНУ ячейку через <br>.\n"
            f"   - Если в переводе встречается символ | внутри текста, замени на &#124;.\n"
            f"   - Пустые ячейки оставляй пустыми (|  |), но только если в оригинале они пустые.\n"
            f"4. ТЕКСТОВЫЕ БЛОКИ: Если в оригинале нет таблицы, не создавай её.\n"
            f"5. ЦИФРЫ, ДАТЫ, ВАЛЮТЫ: Оставляй без изменений. Переводи только текстовые элементы.\n"
            f"6. ВЫВОД: Только переведённый текст. Без вступлений, комментариев, пояснений или оригинала.\n"
            f"\n ПЕРЕД ВЫВОДОМ: Проверь каждую строку таблицы. Количество '|' должно быть одинаковым в заголовке и во всех строках данных."
        )
        return system_prompt, translation_instruction, src_code, tgt_code

    # Точка обработки PDF, изображений, DOCX и TXT
    def process_document(self, file_bytes: bytes, filename: str, target_lang: str = "en", source_lang: Optional[str] = None) -> Dict[str, Any]:
        target_lang = target_lang.lower()
        if target_lang not in ("ru", "en"):
            raise ValueError("target_lang должен быть 'ru' или 'en'")
        
        auto_detect = source_lang is None
        start_time = time.time()
        
        # Временные переменные для кодов (обновятся после определения языка)
        src_code = "??"
        tgt_code = "RU" if target_lang == "ru" else "EN"
        
        logger.info(f"Старт обработки: {filename} | Цель: {tgt_code}")

        try:
            # Поддержка DOCX-файлов
            if filename.lower().endswith(".docx"):
                logger.info("Обработка DOCX-файла")
                original_text = self._docx_to_text(file_bytes)

                if auto_detect:
                    source_lang = self._detect_source_language(original_text[:500])
                    logger.info(f"Автоопределение языка: {source_lang.upper()}")
                else:
                    source_lang = source_lang.lower()

                system_prompt, translation_instruction, src_code, tgt_code = self._get_prompts(source_lang, target_lang)

                chunks = self._split_text(original_text, max_chars=3500)
                translated_chunks = []

                for idx, chunk_text in enumerate(chunks, 1):
                    logger.info(f"DOCX часть {idx}/{len(chunks)}")

                    rag_context = self._build_rag_context(
                        chunk_text[:300],
                        source_lang,
                        target_lang,
                        k=5
                    )

                    text_prompt = (
                        f"{system_prompt}\n\n"
                        f"{rag_context}\n\n"
                        f"ЧАСТЬ ДОКУМЕНТА {idx} ИЗ {len(chunks)}.\n"
                        f"Переведи только эту часть. Не сокращай текст.\n\n"
                        f"ТЕКСТ ДЛЯ ПЕРЕВОДА:\n---\n{chunk_text}\n---\n\n"
                        f"ВЫВЕДИ ТОЛЬКО ПЕРЕВОД:"
                    )

                    response = self.client.generate(
                        model=self.model,
                        prompt=text_prompt,
                        stream=False,
                        options={
                            "temperature": 0.0,
                            "num_predict": 8192,
                            "top_p": 0.9,
                            "repeat_penalty": 1.15
                        }
                    )

                    part_translation = self._parse_response(response)
                    part_translation = self._fix_broken_tables(part_translation, target_lang)
                    translated_chunks.append(part_translation)

                translated_text = "\n\n".join(translated_chunks)
                translated_text = self._fix_broken_tables(translated_text, target_lang)

                elapsed = time.time() - start_time
                pages_count = self._estimate_pages_by_chars(original_text)

                logger.info(
                    f"DOCX готов за {elapsed:.1f}с | "
                    f"{pages_count} стр. | {len(translated_text)} симв."
                )

                return {
                    "success": True,
                    "full_text": translated_text,
                    "text_chunks": [translated_text],
                    "metadata": {
                        "Название файла": filename,
                        "Всего страниц": pages_count,
                        "Всего символов": len(translated_text),
                        "Используемая модель": self.model,
                        "Время обработки, с": round(elapsed, 2),
                        "Среднее время на страницу, с": round(elapsed / pages_count, 2),
                        "Целевой язык": tgt_code,
                        "Исходный язык": source_lang.upper(),
                        "Дата обработки": datetime.now().isoformat(),
                        "RAG включен": self.retriever is not None,
                        "Формат файла": "docx"
                    }
                }
            # Поддержка обычных текстовых файлов
            if filename.lower().endswith(('.txt', '.text', '.md', '.rst')):
                logger.info("Обработка текстового файла (без OCR)")
                original_text = file_bytes.decode('utf-8', errors='replace')
                
                if auto_detect:
                    source_lang = self._detect_source_language(original_text[:500])
                    logger.info(f"Автоопределение языка: {source_lang.upper()}")
                else:
                    source_lang = source_lang.lower()
                
                system_prompt, translation_instruction, src_code, tgt_code = self._get_prompts(source_lang, target_lang)
                rag_context = self._build_rag_context(original_text[:300], source_lang, target_lang, k=5)
                
                text_prompt = (
                    f"{system_prompt}\n\n"
                    f"{rag_context}\n\n"
                    f"ТЕКСТ ДЛЯ ПЕРЕВОДА:\n---\n{original_text}\n---\n\n"
                    f"ВЫВЕДИ ТОЛЬКО ПЕРЕВОД:"
                )
                
                response = self.client.generate(
                    model=self.model, prompt=text_prompt,
                    stream=False, options={
                        "temperature": 0.0, "num_predict": 8192,
                        "top_p": 0.9, "repeat_penalty": 1.15
                    }
                )
                
                translated_text = self._parse_response(response)
                translated_text = self._fix_broken_tables(translated_text, target_lang)
                
                elapsed = time.time() - start_time
                logger.info(f"Готово за {elapsed:.1f}с | {len(translated_text)} симв.")
                pages_count = self._estimate_pages_by_chars(original_text)

                return {
                    "success": True,
                    "full_text": translated_text,
                    "text_chunks": [translated_text],
                    "metadata": {
                        "Название файла": filename,
                        "Всего страниц": pages_count,
                        "Всего символов": len(translated_text),
                        "Используемая модель": self.model,
                        "Время обработки, с": round(elapsed, 2),
                        "Среднее время на страницу, с": round(elapsed / pages_count, 2),
                        "Целевой язык": tgt_code,
                        "Исходный язык": source_lang.upper(),
                        "Дата обработки": datetime.now().isoformat(),
                        "RAG включен": self.retriever is not None,
                        "Формат файла": "text"
                    }
                }
            
            # PDF или изображения
            if filename.lower().endswith(".pdf"):
                images = self._pdf_to_images(file_bytes)
            else:
                img = Image.open(io.BytesIO(file_bytes))
                images = [img.convert("RGB")]

            total_pages = len(images)
            all_text_chunks = []
            previous_text_tail = ""
            detected_source_lang = source_lang
            current_system_prompt = None
            current_translation_instruction = None

            for page_num, image in enumerate(tqdm(images, desc="Страницы", unit="стр"), 1):
                logger.debug(f"Обработка страницы {page_num}/{total_pages}")
                b64_image = self._prepare_image(image)
                
                # Предпросмотр для определения языка
                preview_prompt = "Распознай только первые 2-3 строки текста. Ничего не переводи."
                temp_resp = self.client.generate(
                    model=self.model, prompt=preview_prompt, images=[b64_image],
                    stream=False, options={"temperature": 0.0, "num_predict": 120}
                )
                preview_text = self._parse_response(temp_resp)
                
                # Определяем язык на первой странице
                if auto_detect and page_num == 1:
                    detected_source_lang = self._detect_source_language(preview_text)
                    logger.info(f"Автоопределение языка: {detected_source_lang.upper()}")
                    # ВАЖНО: обновляем промпты ТОЛЬКО после определения языка
                    current_system_prompt, current_translation_instruction, src_code, tgt_code = self._get_prompts(
                        detected_source_lang, target_lang
                    )
                    logger.info(f"Направление перевода: {src_code} → {tgt_code}")
                elif current_system_prompt is None:
                    # Fallback: если по какой-то причине не определили, используем target_lang как подсказку
                    detected_source_lang = "ru" if target_lang == "en" else "en"
                    current_system_prompt, current_translation_instruction, src_code, tgt_code = self._get_prompts(
                        detected_source_lang, target_lang
                    )
                
                # Строим RAG-контекст с правильным направлением
                rag_context = self._build_rag_context(
                    preview_text or previous_text_tail,
                    source_lang=detected_source_lang,
                    target_lang=target_lang,
                    k=3
                )
                
                context_block = ""
                if previous_text_tail:
                    context_block = (
                        f"\nВАЖНО: Продолжение документа. Вот последние строки с предыдущей страницы:\n"
                        f"---\n{previous_text_tail}\n---\n"
                        f"Продолжай перевод в том же стиле. НЕ повторяй заголовки секций.\n"
                    )
                
                page_prompt = (
                    f"[Страница {page_num} из {total_pages}]\n"
                    f"{current_system_prompt}\n"
                    f"{rag_context}\n"
                    f"{context_block}\n"
                    f"\nНапоминание: {current_translation_instruction} Таблицы должны быть идеальными."
                )
                
                response = self.client.generate(
                    model=self.model, prompt=page_prompt, images=[b64_image],
                    stream=False, options={
                        "temperature": 0.0, "num_predict": 4096,
                        "top_p": 0.9, "repeat_penalty": 1.15,
                        "stop": ["--- СТРАНИЦА", "\n\n*Все документы", "END OF DOCUMENT"]
                    }
                )
                
                text = self._parse_response(response)
                if not text:
                    logger.warning(f"Страница {page_num}: пустой ответ")
                    text = f"[Страница {page_num}: не распознано]"
                
                text = self._fix_broken_tables(text, target_lang)
                
                lines = [line for line in text.splitlines() if line.strip()]
                previous_text_tail = "\n".join(lines[-5:]) if lines else ""
                
                chunk = f"--- СТРАНИЦА {page_num} ---\n{text}"
                all_text_chunks.append(chunk)

            elapsed = time.time() - start_time
            full_text = "\n\n".join(all_text_chunks)
            
            logger.info(f"Готово за {elapsed:.1f}с | {total_pages} стр. | {len(full_text)} симв.")

            return {
                "success": True,
                "full_text": full_text,
                "text_chunks": all_text_chunks,
                "metadata": {
                    "Название файла": filename,
                    "Всего страниц": total_pages,
                    "Всего символов": len(full_text),
                    "Используемая модель": self.model,
                    "Время обработки, с": round(elapsed, 2),
                    "Среднее время на страницу, с": round(elapsed / max(total_pages, 1), 2),
                    "Целевой язык": tgt_code,
                    "Исходный язык": detected_source_lang.upper() if detected_source_lang else "UNKNOWN",
                    "Дата обработки": datetime.now().isoformat(),
                    "RAG включен": self.retriever is not None,
                    "Формат файла": "image/pdf"
                }
            }

        except Exception as e:
            logger.error(f"Критическая ошибка: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "full_text": "",
                "text_chunks": [],
                "metadata": {
                    "Название файла": filename,
                    "Время ошибки": datetime.now().isoformat(),
                    "RAG включен": self.retriever is not None,
                    "Целевой язык": target_lang.upper() if 'target_lang' in locals() else "EN"
                }
            }