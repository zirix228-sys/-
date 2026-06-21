import sys
from pathlib import Path
import streamlit as st
import time
import io
import pandas as pd
import csv
import sacrebleu
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt
from datetime import datetime

# Настройка путей для корректного импорта модулей на любой ОС
ROOT_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT_DIR))
from pipeline import DocumentPipeline

# Формирование DOCX-файла из результата перевода
def text_to_docx_bytes(text: str):
    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    lines = text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("|") and line.endswith("|"):
            table_lines = []

            while i < len(lines) and lines[i].strip().startswith("|"):
                current_line = lines[i].strip()
                cells_check = [c.strip() for c in current_line.strip("|").split("|")]
                is_separator = all(set(c) <= set("-: ") for c in cells_check if c)

                if not is_separator:
                    table_lines.append(current_line)

                i += 1

            if table_lines:
                rows = []
                for table_line in table_lines:
                    cells = [cell.strip().replace("<br>", "\n") for cell in table_line.strip("|").split("|")]
                    rows.append(cells)

                max_cols = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=max_cols)
                table.style = "Table Grid"

                for r_idx, row in enumerate(rows):
                    for c_idx in range(max_cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ""
                        table.cell(r_idx, c_idx).text = cell_text

            continue

        if line:
            document.add_paragraph(line)
        else:
            document.add_paragraph("")

        i += 1

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Извлечение текста из эталонного перевода
def extract_text_from_reference(file_bytes: bytes, filename: str) -> str:
    filename = filename.lower()

    if filename.endswith(".docx"):
        document = Document(io.BytesIO(file_bytes))
        parts = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    if filename.endswith((".txt", ".md")):
        return file_bytes.decode("utf-8", errors="replace")

    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(text.strip())

    return "\n".join(parts)

# Расчет BLEU и chrF для сравнения с эталоном
def calculate_translation_metrics(candidate: str, reference: str):
    candidate = candidate.strip()
    reference = reference.strip()

    if not candidate or not reference:
        return None

    bleu = sacrebleu.corpus_bleu([candidate], [[reference]]).score
    chrf = sacrebleu.corpus_chrf([candidate], [[reference]]).score

    return {
        "BLEU": round(bleu, 2),
        "chrF": round(chrf, 2)
    }

# Сохранение истории метрик в CSV и Excel
def save_metrics_history(row: dict):
    results_dir = Path("tests") / "results"
    results_dir.mkdir(parents=True, exist_ok=True)

    csv_path = results_dir / "translation_metrics_history.csv"
    xlsx_path = results_dir / "translation_metrics_history.xlsx"

    file_exists = csv_path.exists()

    with open(csv_path, "a", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=row.keys())

        if not file_exists:
            writer.writeheader()

        writer.writerow(row)

    df = pd.read_csv(csv_path, encoding="utf-8-sig")
    df.to_excel(xlsx_path, index=False)

    return csv_path, xlsx_path

st.set_page_config(
    page_title="Банковский RAG-переводчик",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

if "final_out" not in st.session_state:
    st.session_state.final_out = None

if "uploaded_filename" not in st.session_state:
    st.session_state.uploaded_filename = None

if "metrics_csv_path" not in st.session_state:
    st.session_state.metrics_csv_path = None

if "metrics_xlsx_path" not in st.session_state:
    st.session_state.metrics_xlsx_path = None

if "last_metrics_row" not in st.session_state:
    st.session_state.last_metrics_row = None

st.title("Переводчик банковских документов")
st.caption("Загрузи файл → Qwen-VL + RAG сделают своё дело.")

lang_direction = st.radio(
    "В какую сторону переводим?",
    options=["Русский → English", "English → Русский"],
    index=0,
    horizontal=True
)

lang_map = {
    "Русский → English": "en",
    "English → Русский": "ru"
}
target_lang = lang_map[lang_direction]

uploaded_file = st.file_uploader(
    "Выбери PDF, изображение, текстовый файл или Word-документ",
    type=["pdf", "png", "jpg", "jpeg", "txt", "md", "docx"]
)

# Запуск пайплайна перевода после загрузки файла
if uploaded_file is not None:
    st.success(f"Файл загружен:  `{uploaded_file.name}`")

    if st.button("Запустить перевод", type="primary"):
        file_bytes = uploaded_file.read()
        pipeline = DocumentPipeline()

        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            status_text.text(f"Этап 1/3: Распознавание и перевод ({target_lang.upper()})...")
            progress_bar.progress(33)

            ocr_out = pipeline.run_ocr(file_bytes, uploaded_file.name, target_lang=target_lang)

            status_text.text("Этап 2/3: Форматирование результата...")
            progress_bar.progress(66)
            trans_out = pipeline.run_translate(ocr_out)

            status_text.text("Этап 3/3: Финальная проверка...")
            progress_bar.progress(90)
            final_out = pipeline.run_postprocess(trans_out)

            st.session_state.final_out = final_out
            st.session_state.uploaded_filename = uploaded_file.name

            progress_bar.progress(100)
            status_text.text("Готово")
            time.sleep(0.5)

        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"Ошибка: {str(e)}")
            st.exception(e)


if st.session_state.final_out is not None:
    final_out = st.session_state.final_out
    saved_filename = st.session_state.uploaded_filename or "document"

    st.divider()
    st.subheader("Результат")

    if final_out.get("cleaned_text"):
        st.markdown(final_out["cleaned_text"])
    else:
        st.warning("Текст не извлечён. Убедитесь, что документ читаем.")

    with st.expander("Информация о документе"):
        st.json(final_out.get("metadata", {}))

    if final_out.get("cleaned_text"):
        download_format = st.radio(
            "Формат скачивания:",
            options=["DOCX", "TXT"],
            horizontal=True,
            key="download_format"
        )

        if download_format == "DOCX":
            docx_file = text_to_docx_bytes(final_out["cleaned_text"])

            st.download_button(
                label="Скачать результат (DOCX)",
                data=docx_file.getvalue(),
                file_name=f"translated_{Path(saved_filename).stem}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        else:
            st.download_button(
                label="Скачать результат (TXT)",
                data=final_out["cleaned_text"].encode("utf-8"),
                file_name=f"translated_{Path(saved_filename).stem}.txt",
                mime="text/plain"
            )

    # Сравнение результата перевода с эталонным файлом
    st.divider()
    st.subheader("Сравнение с эталонным переводом")

    translated_text = final_out.get("cleaned_text", "")

    reference_file = st.file_uploader(
        "Загрузи файл с эталонным переводом",
        type=["pdf", "docx", "txt", "md"],
        key="reference_translation_file"
    )

    if reference_file is not None and translated_text:
        if st.button("Рассчитать метрики"):
            reference_bytes = reference_file.read()
            reference_text = extract_text_from_reference(
                reference_bytes,
                reference_file.name
            )

            metrics = calculate_translation_metrics(
                translated_text,
                reference_text
            )

            if metrics:
                metadata = final_out.get("metadata", {})

                row = {
                    "Дата и время": datetime.now().isoformat(),
                    "Исходный файл": st.session_state.uploaded_filename,
                    "Эталонный файл": reference_file.name,
                    "Целевой язык": metadata.get("Целевой язык", metadata.get("target_language", "")),
                    "Исходный язык": metadata.get("Исходный язык", metadata.get("source_language", "")),
                    "RAG включен": metadata.get("RAG включен", metadata.get("rag_enabled", "")),
                    "Модель": metadata.get("Используемая модель", metadata.get("model_used", "")),
                    "Обработано страниц": metadata.get("Всего страниц", metadata.get("pages_processed", "")),
                    "Время обработки, сек": metadata.get("Время обработки, с", metadata.get("processing_time_sec", "")),
                    "Среднее время на страницу, сек": metadata.get("Среднее время на страницу, с", metadata.get("average_time_per_page_sec", "")),
                    "Символов в переводе": len(translated_text),
                    "Символов в эталоне": len(reference_text),
                    "BLEU": metrics["BLEU"],
                    "chrF": metrics["chrF"],
                }

                csv_path, xlsx_path = save_metrics_history(row)
                
                st.session_state.metrics_csv_path = csv_path
                st.session_state.metrics_xlsx_path = xlsx_path
                st.session_state.last_metrics_row = row

                st.success("Метрики рассчитаны и сохранены в историю.")

                st.markdown("""
**BLEU** — метрика от 0 до 100. Она сравнивает машинный перевод с эталонным переводом по совпадению слов и коротких фраз.  
Чем выше BLEU, тем ближе перевод к эталону. Если BLEU близок к 0, часто это значит, что сравниваются разные документы или текст сильно отличается по формулировкам.

**chrF** — метрика от 0 до 100. Она сравнивает тексты на уровне последовательностей символов, а не только целых слов.  
Например, она частично учитывает похожесть словоформ, окончаний, терминов и близких написаний. Поэтому chrF часто мягче и устойчивее BLEU.
""")

            else:
                st.warning("Не удалось рассчитать метрики: перевод или эталонный текст пустой.")
    
        if st.session_state.last_metrics_row is not None:
            st.subheader("Результаты оценки перевода")
            st.json(st.session_state.last_metrics_row)

    if st.session_state.metrics_csv_path is not None and st.session_state.metrics_csv_path.exists():
        with open(st.session_state.metrics_csv_path, "rb") as f:
            st.download_button(
                label="Скачать историю метрик CSV",
                data=f,
                file_name="translation_metrics_history.csv",
                mime="text/csv",
                key="download_metrics_csv"
            )

    if st.session_state.metrics_xlsx_path is not None and st.session_state.metrics_xlsx_path.exists():
        with open(st.session_state.metrics_xlsx_path, "rb") as f:
            st.download_button(
                label="Скачать историю метрик Excel",
                data=f,
                file_name="translation_metrics_history.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="download_metrics_xlsx"
            )